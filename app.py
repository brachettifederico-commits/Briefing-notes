import os
import json
import re
import io
from functools import wraps
from datetime import datetime
from flask import Flask, request, jsonify, send_file, render_template, Response
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

app = Flask(__name__)

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
MODEL = "claude-sonnet-5"

# --- Protezione con password (HTTP Basic Auth) ---
# Impostare su Railway → Variables: APP_USERNAME e APP_PASSWORD
# Se APP_PASSWORD non è impostata, il sito resta accessibile senza password.
APP_USERNAME = os.environ.get("APP_USERNAME", "czp")
APP_PASSWORD = os.environ.get("APP_PASSWORD")


def check_auth(username, password):
    return username == APP_USERNAME and password == APP_PASSWORD


def authenticate():
    return Response(
        "Accesso protetto. Inserisci le credenziali per continuare.",
        401,
        {"WWW-Authenticate": 'Basic realm="Briefing Note Generator"'},
    )


def requires_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not APP_PASSWORD:
            return f(*args, **kwargs)
        auth = request.authorization
        if not auth or not check_auth(auth.username, auth.password):
            return authenticate()
        return f(*args, **kwargs)
    return decorated

SEARCH_PROMPT = """Cerca informazioni aggiornate e dettagliate su questa persona:
- Nome: {nome}{istituzione}
- Focus tematico di interesse: {focus}

Trova e riporta in modo dettagliato:
1. Data e luogo di nascita
2. Ruolo istituzionale attuale
3. Carriera professionale completa (cronologica)
4. Carriera politica completa (cronologica)
5. Almeno 4-6 posizioni, dichiarazioni o attività SPECIFICHE e VERIFICABILI sul tema "{focus}" — non punti generici. Per ognuna cerca una fonte diversa quando possibile (articoli di stampa, comunicati istituzionali, atti parlamentari, interviste).

Regole per le fonti:
- Riporta l'URL ESATTO della pagina che hai effettivamente consultato per ogni fatto, subito dopo il fatto stesso.
- Non riutilizzare lo stesso URL per più di 2 punti diversi: cerca fonti distinte per ciascuna informazione sul focus tematico.
- Se non trovi una fonte affidabile per un punto, ometti il punto piuttosto che inventare una fonte.
- Preferisci fonti primarie (siti istituzionali, comunicati ufficiali, atti parlamentari) e testate giornalistiche riconosciute.

Sii preciso e dettagliato. Includi date, ruoli e contesti specifici."""

JSON_SYSTEM = """Sei un assistente che redige briefing note istituzionali italiane per uno studio di public affairs.

Ti viene fornito un testo con informazioni su una persona, già corredato di URL fonte per ciascun fatto. Struttura queste informazioni in JSON valido.

Regole assolute:
- Rispondi SOLO con JSON valido, nessun testo prima o dopo
- Nessun tag HTML, nessun markdown, nessun <cite>, testo plain
- Italiano formale e istituzionale
- Se un'informazione non è disponibile scrivi "Dato non disponibile"
- Per ogni campo "_url" e ogni "url" nei focus_items: riporta l'URL esatto fornito nel testo di ricerca per quel punto specifico. Se il testo non fornisce un URL per un dato punto, usa null — non inventare e non riutilizzare un URL di un altro punto.
- Nei focus_items, preferisci URL diversi tra loro quando il testo di ricerca li fornisce distinti.

Formato JSON esatto:
{
  "nome": "Nome Cognome",
  "luogo_nascita": "Nato/a a Città (Provincia) il GG mese AAAA",
  "ruolo_attuale": "Ruolo istituzionale attuale completo",
  "background_sintetico": "Ex X, Y e Z (max 10 parole)",
  "istituzione_contesto": "Es: Camera dei Deputati",
  "carriera_professionale": "Testo narrativo prosa cronologico, 150-250 parole, stile istituzionale.",
  "carriera_professionale_url": "URL fonte oppure null",
  "carriera_politica": "Testo narrativo prosa cronologico, 150-250 parole, stile istituzionale.",
  "carriera_politica_url": "URL fonte oppure null",
  "focus_titolo": "Titolo focus tematico (es: Digitale, IA e Cybersecurity)",
  "focus_items": [
    {
      "titolo": "Titolo breve del punto",
      "testo": "Descrizione della posizione o attività. 2-4 frasi plain text.",
      "url": "URL fonte oppure null"
    }
  ]
}

focus_items: 4-6 voci rilevanti al focus tematico richiesto, ciascuna con fonte propria quando disponibile.

ORDINE: ordina i focus_items dal più recente al meno recente, basandoti sulle date/eventi menzionati nel testo di ricerca per ciascun punto (se una voce non ha una data esplicita, posizionala in base al contesto temporale più plausibile). Il primo elemento dell'array deve essere il fatto/dichiarazione più recente."""

# --- Stile grafico del documento Word (formato "Bioprofile" CZP) ---
FONT_NAME = "Trebuchet MS"
NAVY = RGBColor(0x00, 0x20, 0x60)
GRAY_TEXT = RGBColor(0x33, 0x33, 0x33)
GRAY_LINE = "CCCCCC"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ICON_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")
ICON_PROFESSIONALE = os.path.join(ICON_DIR, "icon_professionale.png")


def strip_cite(text):
    """Rimuove tag <cite> e markup dal testo."""
    if not text:
        return text
    text = re.sub(r'<cite[^>]*>([\s\S]*?)</cite>', r'\1', text)
    text = re.sub(r'</?cite[^>]*>', '', text)
    text = re.sub(r'\[?\(Link\)\]?\([^)]*\)', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def clean_obj(obj):
    """Pulisce ricorsivamente tutti i valori stringa da tag HTML."""
    if isinstance(obj, str):
        return strip_cite(obj)
    elif isinstance(obj, list):
        return [clean_obj(i) for i in obj]
    elif isinstance(obj, dict):
        return {k: clean_obj(v) for k, v in obj.items()}
    return obj


def call_anthropic_with_search(prompt):
    """Chiama Anthropic con web search e gestisce il loop tool_use."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    messages = [{"role": "user", "content": prompt}]
    tools = [{"type": "web_search_20250305", "name": "web_search"}]
    text_chunks = []

    for _ in range(10):
        response = client.messages.create(
            model=MODEL,
            max_tokens=8000,
            tools=tools,
            messages=messages,
        )

        messages.append({"role": "assistant", "content": response.content})

        for block in response.content:
            if hasattr(block, "text") and block.text:
                text_chunks.append(block.text)

        if response.stop_reason == "tool_use":
            tool_results = []
            for block in response.content:
                if block.type == "tool_use":
                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": block.id,
                        "content": "Continua la ricerca.",
                    })
            messages.append({"role": "user", "content": tool_results})
            continue

        # end_turn, max_tokens, stop_sequence, ecc.: fermiamo il loop qui.
        # Non richiamare più l'API con l'ultimo messaggio da "assistant",
        # altrimenti l'API rifiuta la richiesta (richiede che l'ultimo
        # messaggio sia "user").
        break

    return "\n".join(text_chunks)


def call_anthropic_json(research_text, focus, data_meeting):
    """Seconda chiamata senza web search per formattare in JSON pulito."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    user_msg = f"""Sulla base di queste informazioni, genera la briefing note in JSON:

INFORMAZIONI RACCOLTE:
{research_text}

FOCUS TEMATICO PER IL CLIENTE: {focus}
{f'DATA MEETING: {data_meeting}' if data_meeting else ''}

Rispondi SOLO con JSON valido, testo plain senza tag HTML."""

    response = client.messages.create(
        model=MODEL,
        max_tokens=4000,
        system=JSON_SYSTEM,
        messages=[{"role": "user", "content": user_msg}],
    )

    for block in response.content:
        if hasattr(block, "text"):
            return block.text

    return None


def add_hyperlink(paragraph, text, url):
    """Aggiunge un hyperlink cliccabile a un paragrafo Word."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_border(cell, sides=("top", "left", "bottom", "right"), val="single", sz="4", color="CCCCCC"):
    """Imposta i bordi di una cella (o li rimuove con val='nil')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side in sides:
        border = tcBorders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            tcBorders.append(border)
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), sz)
        border.set(qn("w:color"), color)


def set_cell_shading(cell, fill="F2F2F2"):
    """Imposta il colore di sfondo di una cella."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_row_height(row, cm_value, rule="exact"):
    """Imposta un'altezza fissa per una riga di tabella."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(cm_value * 567)))
    trHeight.set(qn("w:hRule"), rule)
    trPr.append(trHeight)


def add_paragraph_border(paragraph, sides=("bottom",), color="002060", sz="8", space="8"):
    """Aggiunge una linea sottile sopra/sotto un paragrafo (usata come divisore)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), space)
        border.set(qn("w:color"), color)
        pBdr.append(border)


def set_run_font(run, size=10, color=GRAY_TEXT, bold=False, name=FONT_NAME):
    """Applica font/size/colore/bold a un run in un colpo solo."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    return run


def add_cover_page(doc, d, data_meeting):
    """Copertina navy a piena pagina, stile CZP Bioprofile."""
    if data_meeting:
        try:
            dt = datetime.strptime(data_meeting, "%Y-%m-%d")
            today = dt.strftime("%-d %B %Y")
        except Exception:
            today = data_meeting
    else:
        today = datetime.now().strftime("%-d %B %Y")

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(17)
    set_row_height(table.rows[0], 25.2, rule="exact")
    set_cell_shading(cell, fill="002060")
    set_cell_border(cell, val="nil", color="002060")

    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", "1600"), ("left", "300"), ("right", "300"), ("bottom", "300")]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    p_spacer = cell.paragraphs[0]

    p_line1 = cell.add_paragraph()
    add_paragraph_border(p_line1, sides=("bottom",), color="FFFFFF", sz="6", space="10")

    p_title = cell.add_paragraph()
    r1 = p_title.add_run("BRIEFING")
    set_run_font(r1, size=40, color=WHITE, bold=True)
    r1.add_break()
    r2 = p_title.add_run("NOTE")
    set_run_font(r2, size=40, color=WHITE, bold=True)

    p_sub = cell.add_paragraph()
    sub_text = d.get("istituzione_contesto", "") or "Nota istituzionale"
    r3 = p_sub.add_run(sub_text)
    set_run_font(r3, size=14, color=WHITE, bold=False)
    add_paragraph_border(p_sub, sides=("bottom",), color="FFFFFF", sz="6", space="10")

    p_date = cell.add_paragraph()
    r4 = p_date.add_run(today)
    set_run_font(r4, size=11, color=WHITE, bold=False)

    doc.add_page_break()


def add_circle_photo_placeholder(paragraph, diameter_pt=90):
    """Inserisce un cerchio tratteggiato vuoto (VML) dove incollare la foto a mano."""
    run = paragraph.add_run()
    pict = OxmlElement("w:pict")

    ns = ('xmlns:v="urn:schemas-microsoft-com:vml" '
          'xmlns:w10="urn:schemas-microsoft-com:office:word" '
          'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')
    oval_xml = f'''<v:oval {ns} style="width:{diameter_pt}pt;height:{diameter_pt}pt"
        strokecolor="#999999" strokeweight="1pt" fillcolor="#FAFAFA">
        <v:stroke dashstyle="dash"/>
        <v:textbox inset="0,0,0,0">
            <w:txbxContent>
                <w:p><w:pPr><w:jc w:val="center"/></w:pPr>
                <w:r><w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/><w:color w:val="999999"/><w:sz w:val="16"/></w:rPr><w:t>FOTO</w:t></w:r>
                </w:p>
            </w:txbxContent>
        </v:textbox>
    </v:oval>'''
    oval = parse_xml(oval_xml)
    pict.append(oval)
    run._r.append(pict)


def add_icon_image(paragraph, image_path, width_cm=0.9):
    """Inserisce un'icona immagine (es. omino) centrata in un paragrafo."""
    run = paragraph.add_run()
    run.add_picture(image_path, width=Cm(width_cm))


def add_section(doc, label, body_text=None, url=None, icon_char=None, icon_image=None, is_last=False):
    """Aggiunge una sezione a due colonne (icona/etichetta + contenuto), stile Bioprofile."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    col_label, col_content = table.rows[0].cells
    col_label.width = Cm(3.2)
    col_content.width = Cm(13.8)

    for c in (col_label, col_content):
        set_cell_border(c, val="nil", color="FFFFFF")

    col_label.vertical_alignment = 1
    p_icon = col_label.paragraphs[0]
    p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if icon_image:
        add_icon_image(p_icon, icon_image, width_cm=0.9)
    elif icon_char:
        r_icon = p_icon.add_run(icon_char)
        set_run_font(r_icon, size=18, color=NAVY, bold=False)

    p_label = col_label.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_label = p_label.add_run(label.upper())
    set_run_font(r_label, size=9, color=NAVY, bold=True)

    p = col_content.paragraphs[0]
    r_text = p.add_run(body_text or "")
    set_run_font(r_text, size=10, color=GRAY_TEXT, bold=False)
    if url:
        p.add_run(" ")
        add_hyperlink(p, "(Link)", url)

    if not is_last:
        p_rule = doc.add_paragraph()
        add_paragraph_border(p_rule, sides=("bottom",), color=GRAY_LINE, sz="4", space="4")




def generate_docx(d, data_meeting):
    """Genera il documento Word nel formato CZP (stile Bioprofile)."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    styles = doc.styles
    try:
        styles["Hyperlink"]
    except KeyError:
        hl_style = styles.add_style("Hyperlink", 2)
        hl_style.font.color.rgb = RGBColor(0x11, 0x55, 0xCC)
        hl_style.font.underline = True

    # --- Nome ---
    p = doc.add_paragraph()
    run = p.add_run(d.get("nome", ""))
    set_run_font(run, size=28, color=NAVY, bold=True)
    add_paragraph_border(p, sides=("bottom",), color="002060", sz="12", space="6")

    doc.add_paragraph()

    # --- Riquadro foto tonda + anagrafica ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    photo_cell, info_cell = header_table.rows[0].cells
    photo_cell.width = Cm(3.5)
    info_cell.width = Cm(13.5)
    set_row_height(header_table.rows[0], 3.5, rule="atLeast")

    for c in (photo_cell, info_cell):
        set_cell_border(c, val="nil", color="FFFFFF")
    photo_cell.vertical_alignment = 1
    p_photo = photo_cell.paragraphs[0]
    p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_circle_photo_placeholder(p_photo, diameter_pt=90)

    info_cell.vertical_alignment = 1
    p1 = info_cell.paragraphs[0]
    r_icon = p1.add_run("● ")
    set_run_font(r_icon, size=10, color=NAVY, bold=True)
    r_txt = p1.add_run(d.get("luogo_nascita", ""))
    set_run_font(r_txt, size=10, color=GRAY_TEXT, bold=False)

    p2 = info_cell.add_paragraph()
    r_icon2 = p2.add_run("● ")
    set_run_font(r_icon2, size=10, color=NAVY, bold=True)
    r_txt2 = p2.add_run(d.get("ruolo_attuale", ""))
    set_run_font(r_txt2, size=10, color=GRAY_TEXT, bold=False)

    p_rule = doc.add_paragraph()
    add_paragraph_border(p_rule, sides=("bottom",), color="002060", sz="6", space="6")

    # --- Sezioni: carriera (icona omino) ---
    icon_prof = ICON_PROFESSIONALE if os.path.exists(ICON_PROFESSIONALE) else None

    add_section(doc, "Carriera professionale",
                body_text=d.get("carriera_professionale", ""),
                url=d.get("carriera_professionale_url"),
                icon_image=icon_prof, icon_char=None if icon_prof else "◆")

    add_section(doc, "Carriera politica",
                body_text=d.get("carriera_politica", ""),
                url=d.get("carriera_politica_url"),
                icon_image=icon_prof, icon_char=None if icon_prof else "◆")

    # --- Focus: ogni voce è una sezione a sé con icona lente ---
    focus_items = d.get("focus_items", []) or []
    for i, item in enumerate(focus_items):
        add_section(doc, item.get("titolo", ""),
                    body_text=item.get("testo", ""),
                    url=item.get("url"),
                    icon_char="🔍",
                    is_last=(i == len(focus_items) - 1))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/")
@requires_auth
def index():
    return render_template("index.html")


@app.route("/api/generate", methods=["POST"])
@requires_auth
def generate():
    data = request.get_json()
    nome = data.get("nome", "").strip()
    istituzione = data.get("istituzione", "").strip()
    focus = data.get("focus", "").strip()
    data_meeting = data.get("dataMeeting", "")

    if not nome or not focus:
        return jsonify({"error": "Nome e focus sono obbligatori."}), 400

    try:
        istituzione_str = f"\n- Istituzione / ruolo: {istituzione}" if istituzione else ""
        search_prompt = SEARCH_PROMPT.format(
            nome=nome,
            istituzione=istituzione_str,
            focus=focus
        )
        research_text = call_anthropic_with_search(search_prompt)

        if not research_text:
            return jsonify({"error": "Nessun risultato dalla ricerca."}), 500

        clean_research = strip_cite(research_text)

        json_text = call_anthropic_json(clean_research, focus, data_meeting)
        if not json_text:
            return jsonify({"error": "Nessun JSON nella risposta."}), 500

        clean_json = re.sub(r'```json|```', '', json_text).strip()
        parsed = json.loads(clean_json)
        parsed = clean_obj(parsed)

        return jsonify(parsed)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/docx", methods=["POST"])
@requires_auth
def docx_endpoint():
    data = request.get_json()
    data_meeting = data.get("dataMeeting", "")

    try:
        buf = generate_docx(data, data_meeting)
        nome = data.get("nome", "briefing").replace(" ", "_")
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"Briefing_Note_{nome}.docx"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
